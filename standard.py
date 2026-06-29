import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

# --- 1. CONFIGURATION ---
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

@st.cache_resource
def find_working_model():
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except:
        return "models/gemini-1.5-flash"
    return "models/gemini-1.5-flash"

selected_model_name = find_working_model()
model = genai.GenerativeModel(selected_model_name)

# --- 2. AI LOGIC (INTEGRATED CRITERIA) ---
def generate_advanced_plan(topic, syllabus, extra_context):
    prompt = f"""
    Topic: {topic}. Syllabus Code: {syllabus}. Context: {extra_context}.
    Generate a professional lesson plan in English.
    
    CRITICAL RULES FOR CONTENT FORMATTING:
    1. DO NOT use double asterisks (**) anywhere in your response.
    2. DO NOT use bullet points (e.g., -, *, •) under any circumstances. If a section requires multiple items or a list, you MUST use numbers (1, 2, 3...) exclusively.
    3. All heading markers below must remain in absolute CAPITAL LETTERS.

    Use the following EXACT markers for the document structure:
    
    SECTION: TOPIC
    {topic}
    
    SECTION: LESSON OBJECTIVES
    [Provide exactly 4 numbered points using 1., 2., 3., 4.]
    
    SECTION: LESSON OUTCOMES
    [Provide exactly 4 numbered points using 1., 2., 3., 4.]
    
    SECTION: SUCCESS CRITERIA
    [Provide exactly 4 numbered points using 1., 2., 3., 4.]
    
    SECTION: PREREQUISITE
    [Provide 1 statement]
    
    SECTION: KEYWORDS
    [Provide 6 items separated by commas only. Do not make a list.]
    
    SECTION: HOTS
    1. Analyzing: [Add direct context aligned to the topic]
    2. Evaluating: [Add direct context aligned to the topic]
    3. Creating: [Add direct context aligned to the topic]
    4. Applying: [Add direct context aligned to the topic]
    
    SECTION: DIGITAL CITIZENSHIP
    [Provide exactly 4 numbered points using 1., 2., 3., 4. on ethical tech use/Chromebooks/Canva/YouTube]

    SECTION: OPENING LESSON CONTENT
    [Hook activity and transition plan]

    SECTION: DIFFERENTIATION STRATEGIES (GREEN)
    1. HA (Higher Achiever): [1 challenging activity]

    SECTION: DIFFERENTIATION STRATEGIES (YELLOW)
    1. MA (Medium Achiever): [1 core activity]

    SECTION: DIFFERENTIATION STRATEGIES (RED)
    1. LA (Lower Achiever): [1 scaffolded activity]

    SECTION: BLENDED LEARNING Activity ONE (15 MINS)
    1. Activity 1: [Descriptions]
    2. Teacher Preparation: [Step-by-step before lesson]
    3. Objectives: [3 numbered points]
    4. Student Tasks: [Step-by-step details]

    SECTION: BLENDED LEARNING Activity TWO (15 MINS)
    1. Activity 2: [Descriptions]
    2. Teacher Preparation: [Step-by-step before lesson]
    3. Objectives: [3 numbered points]
    4. Student Tasks: [Step-by-step details]
    
    SECTION: PLENARY (EXIT TICKET)
    [2-3 minute closing activity]

    SECTION: HOMEWORK
    [Task assigned based on topic]

    SECTION: SUGGESTED WAY FORWARD TASK
    [Hook activity and transition plan for next day lesson]
    """
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"System Error: {str(e)}"

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# --- 3. WORD EXPORT LOGIC ---
def create_word_export(topic, syllabus, text):
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_p.add_run()
        header_run.font.name = 'Arial'
        header_run.font.size = Pt(10)
        add_page_number(header_run)

    main_title = f'PTES LESSON PLAN: {topic}'.upper()
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(main_title)
    title_run.font.size = Pt(14)
    title_run.bold = True
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Admin Header Table
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["Week No:", "Date:"], ["Class Size:", "Day:"], ["Venue:", "Duration:"]]
    for r in range(3):
        admin_table.cell(r, 0).text = labels[r][0]
        admin_table.cell(r, 2).text = labels[r][1]
        
    for row in admin_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.size = Pt(12)
    doc.add_paragraph().paragraph_format.line_spacing = 1.0

    # Parsing and Boxing Sections
    sections = text.split('SECTION:')
 
    for section in sections:
        if not section.strip(): continue
        lines = section.strip().split('\n')
        
        title = lines[0].strip().replace("**", "").upper()
        content_lines = lines[1:]
        
        doc_heading = doc.add_paragraph()
        doc_heading.paragraph_format.line_spacing = 1.0
        h_run = doc_heading.add_run(title)
        h_run.bold = True
        h_run.font.size = Pt(14)

        if "KEYWORDS" in title:
            raw_keywords_text = " ".join([l.strip() for l in content_lines if l.strip()])
            keyword_items = [kw.strip() for kw in raw_keywords_text.split(",") if kw.strip()]
            
            kw_table = doc.add_table(rows=2, cols=3)
            kw_table.style = 'Table Grid'
            
            idx = 0
            for r in range(2):
                for c in range(3):
                    if idx < len(keyword_items):
                        cell = kw_table.cell(r, c)
                        cell.text = keyword_items[idx]
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.line_spacing = 1.0
                        if p.runs:
                            p.runs[0].font.size = Pt(12)
                        idx += 1
            doc.add_paragraph().paragraph_format.line_spacing = 1.0
        else:
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            
            # FIXED: Ensures text inside content lines parses accurately for standard blocks including HOTS
            content = "\n".join([l.strip() for l in content_lines if l.strip()]).replace("**", "")
            table.cell(0, 0).text = content
            
            p = table.cell(0, 0).paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            
            # Style reinforcement for table block content
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.line_spacing = 1.0
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
                            
            doc.add_paragraph().paragraph_format.line_spacing = 1.0
     
    # HOD Approval Table Block
    doc.add_page_break()
    
    hod_heading = doc.add_paragraph()
    hod_heading.paragraph_format.line_spacing = 1.0
    hod_run = hod_heading.add_run("HOD APPROVAL & REMARKS")
    hod_run.bold = True
    hod_run.font.size = Pt(14)
    
    hod_table = doc.add_table(rows=2, cols=2)
    hod_table.style = 'Table Grid'
    hod_table.cell(0, 0).text = "Remarks:"
    hod_table.rows[1].height = Pt(50)
    hod_table.cell(1, 0).text = "Date:"
    hod_table.cell(1, 1).text = "Signature:"
    
    for row in hod_table.rows:
        for cell in row.cells:
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            if p.runs:
                p.runs[0].font.size = Pt(12)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 4. GUI ---
st.set_page_config(page_title="Advanced Lesson Planner", layout="wide")

st.title("🎓 PTES Universal Lesson Planner")
st.info("Type in the lesson topic, the subject's syllabus code and the extra information like canva, youtube, infographic")

c1, c2 = st.columns(2)
with c1: u_topic = st.text_input("Lesson Topic:")
with c2: u_syllabus = st.text_input("Syllabus Code:")
u_extra = st.text_area("Extra Context (Optional):")

if st.button("🚀 GENERATE COMPLETE LESSON PLAN"):
    if u_topic and u_syllabus:
        with st.spinner("AI is integrating all criteria into your plan..."):
            result = generate_advanced_plan(u_topic, u_syllabus, u_extra)
            st.session_state['adv_plan_out'] = result.replace("**", "")
    else:
        st.warning("Please fill in the Topic and Syllabus.")

if 'adv_plan_out' in st.session_state:
    st.divider()
    st.subheader("AI Draft Preview")
    st.text_area("Content", st.session_state['adv_plan_out'], height=400)
    doc_file = create_word_export(u_topic, u_syllabus, st.session_state['adv_plan_out'])
    st.download_button("📥 Download to Word version (.docx)", doc_file, f"Universal_LP_{u_topic}.docx")

st.markdown("---")
st.caption("Lesson planner 3.0 | Developer: Hjh Nurul Haziqah Hj Nordin | © 2026 PTES Innovation")
